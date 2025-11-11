from docx import Document
from docx.shared import Inches
import tempfile

# Create a Word document
doc = Document()

# Title
doc.add_heading('Gautam Sharma', 0)
doc.add_paragraph("📍 Hisar, Haryana, India\n📞 +91-XXXXXXXXXX\n📧 gautamsharma@email.com\nLinkedIn | GitHub")

# Summary
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    "Technically driven and passionate IT professional with a strong foundation in cybersecurity, networking, "
    "and programming. Currently employed as a Technical Support Agent at Concentrix, with hands-on experience "
    "in troubleshooting, customer support, and system-level problem-solving. Proficient in tools like Wireshark, "
    "CEH methodologies, and programming in Python, Java, C, and SQL. Eager to build a career in cybersecurity and "
    "networking with continuous upskilling."
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph("Bachelor of Computer Science (B.C.S)\nAsian International University, Manipur\n2021 – 2024 (Expected)")
doc.add_paragraph("Senior Secondary Education (Class 12 – Commerce Stream)\nDelhi Public School (DPS), Hisar\n2020 – 2021")

# Certifications
doc.add_heading('Certifications', level=1)
doc.add_paragraph("- Certified Ethical Hacker (CEH)\n- Cybersecurity Fundamentals (Platform TBD)\n- Networking Basics (e.g., Cisco, CompTIA)")

# Skills
doc.add_heading('Technical Skills', level=1)
doc.add_paragraph(
    "- Cybersecurity: Threat detection, CEH tools, Penetration testing basics\n"
    "- Networking: TCP/IP, DNS, DHCP, Firewalls, OSI model\n"
    "- Programming: Python, Java, C, SQL\n"
    "- Tools: Wireshark, Nmap, Metasploit, Kali Linux\n"
    "- Platforms: Windows, Linux (Ubuntu, Kali)\n"
    "- Other: Ticketing tools, Remote troubleshooting"
)

# Experience
doc.add_heading('Professional Experience', level=1)
doc.add_paragraph(
    "Technical Support Agent\nConcentrix, India\nJune 2024 – Present\n"
    "- Resolved hardware, software, and networking issues\n"
    "- Provided real-time support via call, chat, and remote tools\n"
    "- Logged and tracked support tickets to closure\n"
    "- Delivered customer satisfaction and quick resolution rates"
)

# Projects
doc.add_heading('Projects & Labs', level=1)
doc.add_paragraph(
    "Cybersecurity Lab Project\n- Simulated attacks using Kali Linux\n"
    "- Used Nmap & Wireshark for analysis\n- Practiced vulnerability scanning and patching"
)
doc.add_paragraph(
    "Python Port Scanner\n- Built a basic port scanning tool using sockets\n"
    "- Added logging and threading for efficiency"
)

# Languages
doc.add_heading('Languages', level=1)
doc.add_paragraph("Hindi – Native\nEnglish – Fluent")

# Strengths
doc.add_heading('Strengths', level=1)
doc.add_paragraph(
    "- Excellent communication\n- Strong problem-solving mindset\n"
    "- Eager to learn and grow in cybersecurity field\n- Team player and self-driven"
)

# Save the document
doc.save("Gautam_Sharma_Resume.docx")